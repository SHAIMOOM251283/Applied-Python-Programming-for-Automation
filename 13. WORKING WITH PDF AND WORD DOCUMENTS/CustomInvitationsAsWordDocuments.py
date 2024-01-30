from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_invitations(guest_list, output_file):
    document = Document()

    for guest in guest_list:
        # Create a new paragraph for each line of text
        paragraph1 = document.add_paragraph()
        paragraph2 = document.add_paragraph()
        paragraph3 = document.add_paragraph()
        paragraph4 = document.add_paragraph()
        paragraph5 = document.add_paragraph()

        # Add the invitation text with specified formatting
        run = paragraph1.add_run("It would be a pleasure to have the company of ")
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = 'Palace Script MT'

        run = paragraph2.add_run(f'"{guest}"')
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(255, 0, 0) 
        run.font.name = 'Times New Roman'

        run = paragraph3.add_run(" at 11010 Memory Lane on the evening of ")
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = 'Palace Script MT'

        run = paragraph4.add_run("April 1")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = 'Times New Roman'

        run = paragraph5.add_run(" at 7 o'clock")
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = 'Palace Script MT'

        # Center-align each paragraph
        for paragraph in [paragraph1, paragraph2, paragraph3, paragraph4, paragraph5]:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a page break after each invitation
        document.add_page_break()

    # Save the document to the specified output file
    document.save(output_file)

if __name__ == "__main__":
    # Read the guest names from the file
    with open("guests.txt", "r") as file:
        guest_list = [line.strip() for line in file]

    # Specify the output file for the Word document
    output_file = "invitations.docx"

    # Generate invitations and save to the Word document
    generate_invitations(guest_list, output_file)

    print(f"Invitations generated successfully. Check the '{output_file}' file.")
