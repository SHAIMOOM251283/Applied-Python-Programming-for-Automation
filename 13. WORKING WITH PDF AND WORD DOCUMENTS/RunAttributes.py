#import docx
#doc = docx.Document('demo.docx')
#print(doc.paragraphs[0].text)
#print(doc.paragraphs[0].style)
#doc.paragraphs[0].style = 'Normal'
#print(doc.paragraphs[1].text)
#doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
#print(paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
#doc.paragraphs[1].runs[0].style = 'QuoteChar'
#doc.paragraphs[1].runs[1].underline = True
#doc.paragraphs[1].runs[3].underline = True
#doc.save('restyled.docx')

import docx

doc = docx.Document('demo.docx')

# Print text and style of the first paragraph
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)

# Set style of the first paragraph to 'Normal'
doc.paragraphs[0].style.name = 'Normal'  # Use style name instead of style_id

# Print text of the second paragraph
print(doc.paragraphs[1].text)

# Print text of each run in the second paragraph
for run in doc.paragraphs[1].runs:
    print(run.text)

# Apply styles to runs in the second paragraph
doc.paragraphs[1].runs[0].style.name = 'QuoteChar'  # Use style name instead of style_id
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True

# Save the modified document
doc.save('restyled.docx')


