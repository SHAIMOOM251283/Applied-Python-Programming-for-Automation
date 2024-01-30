#In PyPDF2, when creating a PdfFileReader object, you need to pass a file object instead of a string containing the filename. 
#You can open the PDF file using the built-in open() function and then pass the resulting file object to the PdfFileReader constructor.
import PyPDF2
# Replace 'your_file.pdf' with the actual path to your PDF file
file_path = 'your_file.pdf'
# Open the PDF file in binary mode using the built-in open() function
with open(file_path, 'rb') as file:
    # Pass the file object to PdfFileReader
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Now you can work with the pdf_reader object
    num_pages = pdf_reader.numPages
    # ... perform other operations with the PdfFileReader object

#When working with PdfFileReader and PdfFileWriter in PyPDF2, the file objects need to be opened in binary mode ('rb' for reading and 'wb' for writing). 
#This is because PDF files contain binary data.
#For PdfFileReader (reading a PDF file), you open the file in binary read mode:
with open('your_file.pdf', 'rb') as file:
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Perform operations with pdf_reader
#For PdfFileWriter (writing to a PDF file), you open the file in binary write mode:
with open('output_file.pdf', 'wb') as file:
    pdf_writer = PyPDF2.PdfFileWriter()
    # Perform operations with pdf_writer
    pdf_writer.write(file)

#To acquire a Page object for page 5 from a PdfFileReader object in PyPDF2, you can use the getPage() method. 
#The page indexing is zero-based, so to get the Page object for page 5, you would use index 4.
import PyPDF2
# Replace 'your_file.pdf' with the actual path to your PDF file
file_path = 'your_file.pdf'
# Open the PDF file in binary mode using the built-in open() function
with open(file_path, 'rb') as file:
    # Pass the file object to PdfFileReader
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Get the Page object for page 5 (index 4)
    page_number = 4
    page = pdf_reader.getPage(page_number)
    # Now you can work with the 'page' object
    # For example, you can extract text from the page
    page_text = page.extractText()
    print(page_text)

#The numPages variable in a PdfFileReader object stores the number of pages in the PDF document. You can access it as follows:
import PyPDF2
# Replace 'your_file.pdf' with the actual path to your PDF file
file_path = 'your_file.pdf'
# Open the PDF file in binary mode using the built-in open() function
with open(file_path, 'rb') as file:
    # Pass the file object to PdfFileReader
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Get the number of pages in the PDF document
    num_pages = pdf_reader.numPages
    print(f"The PDF document has {num_pages} pages.")
#In this example, pdf_reader.numPages returns the total number of pages in the PDF document.

#If a PdfFileReader object's PDF is encrypted with a password (e.g., "swordfish"), you need to set the password using the decrypt() method before you can obtain Page objects or perform any other operations on the PDF.
import PyPDF2
# Replace 'your_file.pdf' with the actual path to your encrypted PDF file
file_path = 'your_file.pdf'
# Provide the password
password = 'swordfish'
# Open the PDF file in binary mode using the built-in open() function
with open(file_path, 'rb') as file:
    # Pass the file object to PdfFileReader
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Decrypt the PDF with the password
    if pdf_reader.isEncrypted:
        pdf_reader.decrypt(password)
    # Now you can obtain Page objects or perform other operations
    num_pages = pdf_reader.numPages
    print(f"The PDF document has {num_pages} pages.")
    # Get a specific Page object (e.g., page 0)
    page = pdf_reader.getPage(0)
    # Perform operations with the 'page' object
    page_text = page.extractText()
    print(page_text)
#In this example, pdf_reader.decrypt(password) is used to decrypt the PDF file using the provided password ("swordfish"). After decryption, you can work with the PdfFileReader object as usual, including obtaining Page objects or performing other operations on the PDF.

#In PyPDF2, you can use the rotateClockwise() and rotateCounterClockwise() methods of the PageObject class to rotate a page clockwise or counterclockwise, respectively. 
#These methods allow you to change the orientation of the content on a page.
import PyPDF2
# Replace 'input_file.pdf' and 'output_file.pdf' with the actual paths
input_file_path = 'input_file.pdf'
output_file_path = 'output_file.pdf'
# Open the input PDF file in binary mode
with open(input_file_path, 'rb') as file:
    # Create a PdfFileReader object
    pdf_reader = PyPDF2.PdfFileReader(file)
    # Get a Page object (e.g., page 0)
    page_number = 0
    page = pdf_reader.getPage(page_number)
    # Rotate the page clockwise by 90 degrees
    page.rotateClockwise(90)
    # Create a PdfFileWriter object
    pdf_writer = PyPDF2.PdfFileWriter()
    # Add the rotated page to the writer
    pdf_writer.addPage(page)
    # Write the modified PDF to a new file
    with open(output_file_path, 'wb') as output_file:
        pdf_writer.write(output_file)

#In Python's python-docx library, you can use the Document class to work with Word documents. 
#To get a Document object for a file named "demo.docx," you can use the Document() constructor and pass the file path as an argument.
from docx import Document
# Replace 'demo.docx' with the actual path to your Word document
file_path = 'demo.docx'
# Create a Document object for the Word document
doc = Document(file_path)
# Now you can work with the 'doc' object, which represents the Word document
# For example, you can iterate through paragraphs in the document
for paragraph in doc.paragraphs:
    print(paragraph.text)

#In the context of the python-docx library for working with Word documents in Python, a Paragraph object and a Run object represent different elements of the document structure.
#Paragraph Object:
    #A Paragraph object represents a paragraph in a Word document.
    #A paragraph is a block of text that ends with a hard return.
    #Paragraphs can contain multiple runs and other elements like tables, images, and more.
    #You can access paragraphs in a document using the paragraphs attribute of the Document object.
#Example:
from docx import Document
doc = Document('example.docx')
# Accessing paragraphs
for paragraph in doc.paragraphs:
    print(paragraph.text)
#Run Object:
    #A Run object represents a contiguous run of text with the same style in a paragraph.
    #Within a paragraph, text can have different styles such as bold, italic, font size, etc. Each of these styled portions is represented by a Run object.
    #You can access runs within a paragraph using the runs attribute of the Paragraph object.
#Example:
from docx import Document
doc = Document('example.docx')
# Accessing runs within a paragraph
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        print(run.text)
#In summary, a Paragraph object represents a block of text in a Word document, while a Run object represents a styled portion of text within a paragraph. 
#Understanding and working with both Paragraph and Run objects is essential when manipulating the content and formatting of Word documents using the python-docx library.

#To obtain a list of Paragraph objects for a Document object stored in a variable named doc, you can access the paragraphs attribute of the Document object
from docx import Document
# Assuming 'doc' is a Document object (replace with the actual document variable)
# For example, you can create a Document object from a file named 'example.docx'
doc = Document('example.docx')
# Accessing the list of Paragraph objects
paragraphs_list = doc.paragraphs
# Now you can iterate through the list of paragraphs and work with each Paragraph object
for paragraph in paragraphs_list:
    print(paragraph.text)
#In this example, doc.paragraphs returns a list of Paragraph objects from the Document object. 
#You can then iterate through this list to access and work with each individual Paragraph object. 
#Each Paragraph object has attributes and methods that allow you to manipulate and retrieve information about the text and formatting within that paragraph.

#In the context of the python-docx library for working with Word documents in Python, the object that has attributes such as bold, underline, italic, strike, and outline is a Run object.
#A Run object represents a contiguous run of text with the same style within a paragraph. 
#These style attributes can include formatting options like bold, underline, italic, strike-through, outline, font size, and more.
from docx import Document
# Assuming 'doc' is a Document object (replace with the actual document variable)
# For example, you can create a Document object from a file named 'example.docx'
doc = Document('example.docx')
# Accessing runs and their attributes within a paragraph
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        print(f"Text: {run.text}")
        print(f"Bold: {run.bold}")
        print(f"Underline: {run.underline}")
        print(f"Italic: {run.italic}")
        print(f"Strike: {run.strike}")
        print(f"Outline: {run.outline}")
        print()

#In the context of the python-docx library for working with Word documents in Python:
#Setting bold to True:
    #If you set the bold attribute of a Run object to True, it means that the text represented by that Run will be formatted as bold.
run.bold = True
#Setting bold to False:
    #If you set the bold attribute of a Run object to False, it means that the text represented by that Run will not be formatted as bold.
run.bold = False
#Setting bold to None:
    #If you set the bold attribute of a Run object to None, it means that the bold formatting will follow the default style defined by the underlying style hierarchy. 
    #It inherits the style from the surrounding context (paragraph style, document style, etc.).
run.bold = None
#In summary:
    #True: Explicitly sets the text to bold.
    #False: Explicitly sets the text to not be bold.
    #None: Inherits the bold formatting from the surrounding context.

#To create a Document object for a new Word document using the python-docx library in Python, you can simply use the Document constructor without specifying a file path. 
#This creates an in-memory document that you can then manipulate and save to a file if needed.
from docx import Document
# Create a new Document object
doc = Document()
# Now you can add content to the document
doc.add_heading('My New Document', level=1)
doc.add_paragraph('This is a new Word document.')
# Save the document to a file (optional)
doc.save('new_document.docx')
#In this example:
    #Document() creates a new Document object.
    #add_heading() adds a heading to the document.
    #add_paragraph() adds a paragraph to the document.
#If you want to save the newly created document to a file, you can use the save() method and provide the desired file name (e.g., 'new_document.docx'). 
#If you don't save it explicitly, the document remains in-memory and can be further modified or used as needed.

#To add a paragraph with the text 'Hello there!' to a Document object stored in a variable named doc using the python-docx library, you can use the add_paragraph method.
from docx import Document
# Assuming 'doc' is a Document object (replace with the actual document variable)
# For example, you can create a new Document object using `Document()`
doc = Document()
# Add a paragraph with the text 'Hello there!'
doc.add_paragraph('Hello there!')
# Now you can continue working with the 'doc' object
# Save the document to a file (optional)
doc.save('example.docx')
#In this example, doc.add_paragraph('Hello there!') adds a new paragraph with the specified text to the Document object. 
#If you've already created a Document object and stored it in the variable doc, you can use the add_paragraph method to add content to it.

#In Word documents and in the python-docx library, headings are often organized into different levels, and each level is assigned an integer value. 
#The integer values represent the hierarchy of the headings, with lower numbers indicating higher-level headings. 
#In the context of Word documents and python-docx, the commonly used heading levels and their corresponding integer values are:
    #Heading 1: Level 1
    #Heading 2: Level 2
    #Heading 3: Level 3
    #Heading 4: Level 4
    #Heading 5: Level 5
    #Heading 6: Level 6
#Here are the corresponding integer values for each heading level:
    #Heading 1: Level 1 (integer value: 0)
    #Heading 2: Level 2 (integer value: 1)
    #Heading 3: Level 3 (integer value: 2)
    #Heading 4: Level 4 (integer value: 3)
    #Heading 5: Level 5 (integer value: 4)
    #Heading 6: Level 6 (integer value: 5)
from docx import Document
doc = Document()
# Adding Heading 1
doc.add_heading('Heading 1 Text', level=1)
# Save the document to a file (optional)
doc.save('example.docx')
#In this example, the level=1 parameter indicates that 'Heading 1 Text' should be formatted as a Heading 1 in the Word document. 
#Adjust the level parameter accordingly for other heading levels.

