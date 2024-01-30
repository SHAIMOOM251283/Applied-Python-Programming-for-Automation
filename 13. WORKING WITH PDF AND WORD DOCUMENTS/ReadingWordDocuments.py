import docx
# Load the Word document
doc = docx.Document('demo.docx')
# Print the total number of paragraphs in the document
print("Total number of paragraphs:", len(doc.paragraphs))
# Print the text of the first paragraph
print("Text of the first paragraph:", doc.paragraphs[0].text)
# Print the text of the second paragraph
print("Text of the second paragraph:", doc.paragraphs[1].text)
# Print the total number of runs in the second paragraph
print("Total number of runs in the second paragraph:", len(doc.paragraphs[1].runs))
# Print the text of each run in the second paragraph
for index, run in enumerate(doc.paragraphs[1].runs):
    print(f"Run {index + 1}: {run.text}")

#import docx
# Load the Word document
#doc = docx.Document('demo.docx')
# Print the total number of paragraphs in the document
#print("Total number of paragraphs:", len(doc.paragraphs))
# Print the text of each paragraph
#for index, paragraph in enumerate(doc.paragraphs):
#    print(f"Paragraph {index + 1}:", paragraph.text)
# Print details about runs in the second paragraph
#second_paragraph_runs = doc.paragraphs[1].runs
#print("Total number of runs in the second paragraph:", len(second_paragraph_runs))
# Print the text of each run in the second paragraph
#for index, run in enumerate(second_paragraph_runs):
#    print(f"Run {index + 1}: {run.text}")







