#import docx
#doc = docx.Document()
#doc.add_paragraph('This is on the first page!')
#doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
#doc.add_paragraph('This is on the second page!')
#doc.save('twoPage.docx')

import docx
from docx.enum.text import WD_BREAK
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')
