#! python3
#import docx, readDocx
#def getText(filename):
#   doc = docx.Document(filename)
#   fullText = []
#   for para in doc.paragraphs:
#    fullText.append(para.text)
#    return '\n'.join(fullText)
#print(readDocx.getText('demo.docx'))

#import docx
#def get_text(filename):
#    doc = docx.Document(filename)
#    full_text = []
#    for para in doc.paragraphs:
#        full_text.append(para.text)
#    return '\n'.join(full_text)
#if __name__ == "__main__":
    # Example usage
#    document_path = 'demo.docx'
#    text_content = get_text(document_path)
#    print(text_content)

#import docx
#def get_text(filename):
#    doc = docx.Document(filename)
#    full_text = []
#    for para in doc.paragraphs:
#        full_text.append(' ' + para.text)
#    return '\n'.join(full_text)
#if __name__ == "__main__":
    # Example usage
#    document_path = 'demo.docx'
#    text_content = get_text(document_path)
#    print(text_content)


import docx

def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Add two newlines between each paragraph
    return '\n\n'.join(full_text)


if __name__ == "__main__":
    # Example usage
    document_path = 'demo.docx'
    text_content = get_text(document_path)
    print(text_content)


